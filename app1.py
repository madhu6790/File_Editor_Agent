import streamlit as st
from docx import Document
import pandas as pd
from PyPDF2 import PdfReader
import io
import requests
import openpyxl
from dotenv import load_dotenv
import os

# Load environment variables from .env file
load_dotenv()
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_API_URL = "https://api.groq.com/openai/v1/chat/completions"

def extract_text(file, file_type):
    if file_type == "docx":
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    elif file_type == "pdf":
        reader = PdfReader(file)
        return "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
    elif file_type == "txt":
        return file.read().decode("utf-8")
    elif file_type == "csv":
        df = pd.read_csv(file)
        return df.to_csv(index=False)
    elif file_type == "xlsx":
        df = pd.read_excel(file)
        return df.to_csv(index=False)
    else:
        return ""

def create_file(content, file_type):
    output = io.BytesIO()

    if file_type == "docx":
        doc = Document()
        for line in content.split("\n"):
            doc.add_paragraph(line)
        doc.save(output)

    elif file_type == "pdf":
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
        c = canvas.Canvas(output, pagesize=letter)
        text_object = c.beginText(1 * inch, 10.5 * inch)
        text_object.setFont("Arial", 10)

        for line in content.split("\n"):
            text_object.textLine(line)
        c.drawText(text_object)
        c.save()

    elif file_type in ["csv", "xlsx"]:
        from io import StringIO
        df = pd.read_csv(StringIO(content))
        if file_type == "csv":
            df.to_csv(output, index=False)
        else:
            df.to_excel(output, index=False, engine="openpyxl")

    elif file_type == "txt":
        output.write(content.encode("utf-8"))

    output.seek(0)
    return output

def call_groq_editor(text, prompt):
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    data = {
        "model": "llama3-8b-8192",
        "messages": [
            {"role": "system", "content": "You are a helpful editor. Edit the text as per user instructions."},
            {"role": "user", "content": f"Here is the text:\n{text}\n\nPlease edit it according to this instruction:\n{prompt}\n\nReturn the full edited text."}
        ]
    }
    response = requests.post(GROQ_API_URL, headers=headers, json=data)
    response.raise_for_status()
    return response.json()["choices"][0]["message"]["content"]

def main():
    st.title("File Editor Agent")

    uploaded_file = st.file_uploader("Upload a file", type=["docx", "pdf", "txt", "csv", "xlsx"])
    prompt = st.text_area(" Enter editing instructions", "Replace 'old_Text' with 'New_Text'")

    if uploaded_file and prompt:
        file_name = uploaded_file.name
        file_type = file_name.split(".")[-1].lower()

        if st.button("✨ Edit File"):
            with st.spinner("Editing in progress..."):

                # For CSV/XLSX use DataFrame round-trip
                if file_type in ["csv", "xlsx"]:
                    df = pd.read_csv(uploaded_file) if file_type == "csv" else pd.read_excel(uploaded_file)
                    original_text = df.to_csv(index=False)
                    edited_text = call_groq_editor(original_text, prompt)

                    from io import StringIO
                    edited_df = pd.read_csv(StringIO(edited_text))

                    output = io.BytesIO()
                    if file_type == "csv":
                        edited_df.to_csv(output, index=False)
                    else:
                        edited_df.to_excel(output, index=False, engine="openpyxl")
                    output.seek(0)

                    st.success(" Editing complete!")
                    st.download_button(
                        label=f" Download Edited File ({file_type})",
                        data=output,
                        file_name=f"edited_file.{file_type}",
                        mime="application/octet-stream"
                    )

                else:
                    original_text = extract_text(uploaded_file, file_type)
                    edited_text = call_groq_editor(original_text, prompt)
                    edited_stream = create_file(edited_text, file_type)

                    st.success(" Editing complete!")
                    st.download_button(
                        label=f" Download Edited File ({file_type})",
                        data=edited_stream,
                        file_name=f"edited_file.{file_type}",
                        mime="application/octet-stream"
                    )

if __name__ == "__main__":
    main()
